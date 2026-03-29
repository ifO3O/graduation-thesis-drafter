#!/usr/bin/env python3
"""
Render thesis markdown into school-formatted DOCX using a template.
"""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.text.paragraph import Paragraph

CHAPTER_RE = re.compile(r"^\s*第[0-9一二三四五六七八九十百千]+\s*章")
H2_RE = re.compile(r"^\s*\d+\.\d+\s+")
H3_RE = re.compile(r"^\s*\d+\.\d+\.\d+\s+")


def ensure_pandoc(override_path: str | None = None) -> str:
    if override_path:
        p = Path(override_path).resolve()
        if p.exists():
            return str(p)
        raise SystemExit(f"pandoc path not found: {p}")

    pandoc = shutil.which("pandoc")
    if pandoc:
        return pandoc

    local_app = os.environ.get("LOCALAPPDATA", "")
    candidates = [
        Path(local_app) / "Pandoc" / "pandoc.exe",
        Path(local_app) / "Microsoft" / "WinGet" / "Links" / "pandoc.exe",
        Path("C:/Program Files/Pandoc/pandoc.exe"),
        Path("C:/Program Files (x86)/Pandoc/pandoc.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)

    raise SystemExit(
        "pandoc not found. Install pandoc or pass --pandoc <path-to-pandoc>."
    )


def load_spec_from_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def build_spec_with_extractor(template_docx: Path) -> dict[str, Any]:
    this_dir = Path(__file__).resolve().parent
    extractor = this_dir / "extract_docx_style_spec.py"
    if not extractor.exists():
        raise SystemExit(f"Missing script: {extractor}")

    spec_module = importlib.util.spec_from_file_location("extract_docx_style_spec", extractor)
    if spec_module is None or spec_module.loader is None:
        raise SystemExit("Cannot load extract_docx_style_spec.py")

    module = importlib.util.module_from_spec(spec_module)
    spec_module.loader.exec_module(module)

    doc = Document(str(template_docx))
    paragraphs = list(doc.paragraphs)
    non_empty = [p for p in paragraphs if p.text.strip()]

    style_map = module.detect_style_map(non_empty)
    toc_styles = module.collect_toc_styles(non_empty)
    chapter_paras = [p for p in non_empty if style_map["chapter"] and p.style.name == style_map["chapter"]]
    h2_paras = [p for p in non_empty if style_map["heading2"] and p.style.name == style_map["heading2"]]
    h3_paras = [p for p in non_empty if style_map["heading3"] and p.style.name == style_map["heading3"]]

    body_candidates = []
    for p in non_empty:
        if p.style.name in {style_map["chapter"], style_map["heading2"], style_map["heading3"]}:
            continue
        if len(p.text.strip()) >= 15:
            body_candidates.append(p)
    if not body_candidates:
        body_candidates = non_empty

    return {
        "template_file": str(template_docx),
        "page_setup": module.build_page_setup(doc),
        "style_map": style_map,
        "toc_styles": toc_styles,
        "format_rules": {
            "body": module.choose_majority_signature(body_candidates),
            "chapter": module.choose_majority_signature(chapter_paras),
            "heading2": module.choose_majority_signature(h2_paras),
            "heading3": module.choose_majority_signature(h3_paras),
        },
    }


def parse_bool(value: Any) -> bool | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.lower() == "true"
    return bool(value)


def set_run_fonts(run, font_info: dict[str, Any], size_pt: float | None, bold: bool | None) -> None:
    east_asia = font_info.get("eastAsia")
    ascii_font = font_info.get("ascii")
    hansi = font_info.get("hAnsi")

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()

    if east_asia:
        r_fonts.set(qn("w:eastAsia"), str(east_asia))
    if ascii_font:
        r_fonts.set(qn("w:ascii"), str(ascii_font))
        run.font.name = str(ascii_font)
    if hansi:
        r_fonts.set(qn("w:hAnsi"), str(hansi))

    if size_pt is not None:
        run.font.size = Pt(float(size_pt))
    if bold is not None:
        run.bold = bool(bold)


def apply_rule_to_paragraph(paragraph, rule: dict[str, Any]) -> None:
    if not rule:
        return

    p_fmt = paragraph.paragraph_format
    line_spacing = rule.get("line_spacing")
    first_indent_pt = rule.get("first_line_indent_pt")
    alignment = rule.get("alignment")

    if line_spacing not in (None, "None", ""):
        try:
            line_value = float(line_spacing)
            p_fmt.line_spacing = line_value
        except Exception:
            pass

    if first_indent_pt not in (None, "None", ""):
        try:
            p_fmt.first_line_indent = Pt(float(first_indent_pt))
        except Exception:
            pass

    if isinstance(alignment, str):
        if "LEFT" in alignment:
            paragraph.alignment = 0
        elif "CENTER" in alignment:
            paragraph.alignment = 1
        elif "RIGHT" in alignment:
            paragraph.alignment = 2
        elif "JUSTIFY" in alignment:
            paragraph.alignment = 3

    font_info = rule.get("font", {}) if isinstance(rule.get("font"), dict) else {}
    size_pt = rule.get("size_pt")
    bold = parse_bool(rule.get("bold"))

    for run in paragraph.runs:
        if run.text and run.text.strip():
            set_run_fonts(run, font_info, size_pt, bold)


def classify_paragraph(paragraph) -> str:
    text = paragraph.text.strip()
    style_name = (paragraph.style.name or "").strip().lower()

    if style_name in {"heading 1", "heading1", "标题 1"}:
        return "chapter"
    if style_name in {"heading 2", "heading2", "标题 2"}:
        return "heading2"
    if style_name in {"heading 3", "heading3", "标题 3"}:
        return "heading3"

    if CHAPTER_RE.search(text) or text in {"摘 要", "ABSTRACT", "目录", "目 录", "参考文献", "致谢", "致 谢"}:
        return "chapter"
    if H3_RE.search(text):
        return "heading3"
    if H2_RE.search(text):
        return "heading2"
    return "body"


def apply_page_setup(doc: Document, page_setup: dict[str, Any]) -> None:
    for section in doc.sections:
        if page_setup.get("page_width_mm") is not None:
            section.page_width = Mm(float(page_setup["page_width_mm"]))
        if page_setup.get("page_height_mm") is not None:
            section.page_height = Mm(float(page_setup["page_height_mm"]))
        if page_setup.get("top_margin_mm") is not None:
            section.top_margin = Mm(float(page_setup["top_margin_mm"]))
        if page_setup.get("bottom_margin_mm") is not None:
            section.bottom_margin = Mm(float(page_setup["bottom_margin_mm"]))
        if page_setup.get("left_margin_mm") is not None:
            section.left_margin = Mm(float(page_setup["left_margin_mm"]))
        if page_setup.get("right_margin_mm") is not None:
            section.right_margin = Mm(float(page_setup["right_margin_mm"]))
        if page_setup.get("gutter_mm") is not None:
            section.gutter = Mm(float(page_setup["gutter_mm"]))
        if page_setup.get("header_distance_mm") is not None:
            section.header_distance = Mm(float(page_setup["header_distance_mm"]))
        if page_setup.get("footer_distance_mm") is not None:
            section.footer_distance = Mm(float(page_setup["footer_distance_mm"]))


def find_toc_heading_index(paragraphs: list[Any]) -> int | None:
    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if txt in {"目 录", "目录", "Table of Contents"}:
            return i
    return None


def paragraph_has_toc_field(paragraph) -> bool:
    fields = paragraph._p.xpath('.//w:fldSimple[contains(@w:instr, "TOC")]')
    return len(fields) > 0


def insert_paragraph_after(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_toc_field(paragraph) -> None:
    if paragraph_has_toc_field(paragraph):
        return

    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')

    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "目录（在 Word 中右键目录并更新域）"
    run.append(text)
    fld_simple.append(run)

    paragraph._p.append(fld_simple)


def ensure_toc_placeholder(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    if any(paragraph_has_toc_field(p) for p in paragraphs):
        return

    idx = find_toc_heading_index(paragraphs)
    if idx is None:
        return

    heading_para = paragraphs[idx]
    target = insert_paragraph_after(heading_para)
    add_toc_field(target)


def apply_spec_to_docx(docx_path: Path, spec: dict[str, Any]) -> None:
    doc = Document(str(docx_path))
    apply_page_setup(doc, spec.get("page_setup", {}))

    style_map = spec.get("style_map", {})
    format_rules = spec.get("format_rules", {})

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        category = classify_paragraph(para)

        style_name = style_map.get(category)
        if style_name:
            try:
                para.style = style_name
            except Exception:
                pass

        apply_rule_to_paragraph(para, format_rules.get(category, {}))

    ensure_toc_placeholder(doc)
    doc.save(str(docx_path))


def run_pandoc(
    markdown: Path,
    template_docx: Path,
    output_docx: Path,
    pandoc_path: str | None,
    number_sections: bool,
) -> None:
    pandoc = ensure_pandoc(pandoc_path)
    cmd = [
        pandoc,
        str(markdown),
        "-o",
        str(output_docx),
        "--reference-doc",
        str(template_docx),
        "--toc",
        "--toc-depth=3",
    ]
    if number_sections:
        cmd.append("--number-sections")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise SystemExit(f"pandoc failed:\n{result.stderr.strip()}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Render thesis markdown using school DOCX template.")
    parser.add_argument("--markdown", required=True, help="Input markdown thesis draft path")
    parser.add_argument("--template-docx", required=True, help="School thesis template DOCX")
    parser.add_argument("--output-docx", required=True, help="Output DOCX path")
    parser.add_argument("--style-spec-json", help="Optional pre-generated style spec JSON")
    parser.add_argument("--cover-docx", help="Optional cover DOCX to copy beside output")
    parser.add_argument("--pandoc", help="Optional explicit path to pandoc executable")
    parser.add_argument(
        "--number-sections",
        action="store_true",
        help="Enable pandoc automatic section numbering",
    )
    args = parser.parse_args()

    markdown = Path(args.markdown).resolve()
    template_docx = Path(args.template_docx).resolve()
    output_docx = Path(args.output_docx).resolve()

    if not markdown.exists():
        raise SystemExit(f"Markdown not found: {markdown}")
    if not template_docx.exists():
        raise SystemExit(f"Template DOCX not found: {template_docx}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    run_pandoc(markdown, template_docx, output_docx, args.pandoc, args.number_sections)

    if args.style_spec_json:
        spec = load_spec_from_json(Path(args.style_spec_json).resolve())
    else:
        spec = build_spec_with_extractor(template_docx)

    apply_spec_to_docx(output_docx, spec)
    print(f"Rendered thesis DOCX: {output_docx}")

    if args.cover_docx:
        cover = Path(args.cover_docx).resolve()
        if cover.exists():
            cover_out = output_docx.with_name(output_docx.stem + "_相关资料封面.docx")
            try:
                shutil.copy2(cover, cover_out)
                print(f"Copied cover DOCX: {cover_out}")
            except Exception as exc:
                print(f"Cover copy failed, skipped: {exc}")
        else:
            print(f"Cover DOCX not found, skipped: {cover}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
